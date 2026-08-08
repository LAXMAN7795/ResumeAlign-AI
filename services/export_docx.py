import io
import re
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

def markdown_to_docx_bytes(markdown_text: str) -> bytes:
    """Converts a Markdown resume string into an editable DOCX file buffer using python-docx."""
    doc = Document()

    # Page Margins (0.6 inch margins for modern layout)
    sections = doc.sections
    for section in sections:
        section.top_margin = Inches(0.6)
        section.bottom_margin = Inches(0.6)
        section.left_margin = Inches(0.6)
        section.right_margin = Inches(0.6)

    lines = markdown_text.split('\n')

    for line in lines:
        stripped = line.strip()
        if not stripped:
            continue

        # Document Title (# Name)
        if stripped.startswith('# '):
            text = stripped[2:].strip()
            p = doc.add_paragraph()
            run = p.add_run(text)
            run.font.name = 'Calibri'
            run.font.size = Pt(18)
            run.font.bold = True
            run.font.color.rgb = RGBColor(15, 23, 42)
            p.paragraph_format.space_after = Pt(6)

        # Section Heading (## Section)
        elif stripped.startswith('## '):
            text = stripped[3:].strip()
            p = doc.add_paragraph()
            run = p.add_run(text.upper())
            run.font.name = 'Calibri'
            run.font.size = Pt(12)
            run.font.bold = True
            run.font.color.rgb = RGBColor(37, 99, 235)
            p.paragraph_format.space_before = Pt(10)
            p.paragraph_format.space_after = Pt(3)

        # Sub-heading (### Role / Project)
        elif stripped.startswith('### '):
            text = stripped[4:].strip()
            p = doc.add_paragraph()
            run = p.add_run(text)
            run.font.name = 'Calibri'
            run.font.size = Pt(11)
            run.font.bold = True
            p.paragraph_format.space_before = Pt(4)
            p.paragraph_format.space_after = Pt(2)

        # Bullet List Items
        elif stripped.startswith(('* ', '- ', '• ')):
            bullet_text = stripped[2:].strip()
            p = doc.add_paragraph(style='List Bullet')
            p.paragraph_format.space_after = Pt(2)
            
            # Helper to parse **bold** inside bullets
            parts = re.split(r'(\*\*.*?\*\*)', bullet_text)
            for part in parts:
                if part.startswith('**') and part.endswith('**'):
                    run = p.add_run(part[2:-2])
                    run.bold = True
                else:
                    run = p.add_run(part)
                run.font.name = 'Calibri'
                run.font.size = Pt(10)

        # Standard Text Paragraphs
        else:
            p = doc.add_paragraph()
            p.paragraph_format.space_after = Pt(4)
            parts = re.split(r'(\*\*.*?\*\*)', stripped)
            for part in parts:
                if part.startswith('**') and part.endswith('**'):
                    run = p.add_run(part[2:-2])
                    run.bold = True
                else:
                    run = p.add_run(part)
                run.font.name = 'Calibri'
                run.font.size = Pt(10)

    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer.getvalue()